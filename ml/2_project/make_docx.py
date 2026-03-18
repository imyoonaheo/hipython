"""
DR 시뮬레이션 대시보드 사용 설명서 Word 문서 생성 스크립트
python-docx 사용
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── 페이지 여백 설정 ───────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = int(8.5 * 914400)   # US Letter
section.page_height = int(11  * 914400)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ─── 스타일 헬퍼 함수 ───────────────────────────────────────────────────────
def set_run_style(run, bold=False, size=12, color=None, font='맑은 고딕'):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if color:
        run.font.color.rgb = RGBColor(*color)

def set_cell_bg(cell, hex_color):
    """셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color='CCCCCC'):
    """셀 테두리 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(text, level=1, color=(0x1e, 0x40, 0xaf)):
    """제목 단락 추가"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    if level == 1:
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '3B82F6')
        pBdr.append(bottom)
        pPr.append(pBdr)
    run = p.add_run(text)
    sizes = {1: 18, 2: 14, 3: 12}
    set_run_style(run, bold=True, size=sizes.get(level, 12), color=color)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body(text, indent=0, bullet=False, color=None):
    """본문 단락 추가"""
    p = doc.add_paragraph()
    if bullet:
        p.style = 'List Bullet'
    if indent:
        p.paragraph_format.left_indent = Inches(indent * 0.3)
    run = p.add_run(text)
    set_run_style(run, size=10.5, color=color)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_note_box(text, bg='EFF6FF', border_color='2563EB'):
    """안내 박스 (표로 구현)"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    set_run_style(run, size=10, color=(0x1e, 0x40, 0xaf))
    cell.paragraphs[0].paragraph_format.space_before = Pt(4)
    cell.paragraphs[0].paragraph_format.space_after = Pt(4)
    cell.paragraphs[0].paragraph_format.left_indent = Inches(0.1)
    # 테두리 색 변경
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), border_color)
    tcBorders.append(left)
    tcPr.append(tcBorders)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 표지 / 헤더
# ═══════════════════════════════════════════════════════════════════════════════
# 상단 색상 배너 (표로 구현)
banner = doc.add_table(rows=1, cols=1)
banner.alignment = WD_TABLE_ALIGNMENT.CENTER
banner_cell = banner.cell(0, 0)
set_cell_bg(banner_cell, '1E40AF')  # 진한 파란색

for i in range(3):
    if i == 0:
        p = banner_cell.paragraphs[0]
    else:
        p = banner_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

title_lines = [
    ('⚡ DR 시뮬레이션 대시보드', 22, True, (255, 255, 255)),
    ('사용 설명서', 18, False, (186, 230, 253)),
    ('수요반응(DR) 발령 시 예상 이익 계산 도구  |  제조 공장 에너지 담당자용', 10, False, (186, 230, 253)),
]
for idx, (text, sz, bold, col) in enumerate(title_lines):
    if idx == 0:
        p = banner_cell.paragraphs[0]
    else:
        p = banner_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_style(run, bold=bold, size=sz, color=col)
    p.paragraph_format.space_before = Pt(8 if idx == 0 else 4)
    p.paragraph_format.space_after = Pt(8 if idx == 2 else 2)

doc.add_paragraph()

# 메타 정보
meta_tbl = doc.add_table(rows=1, cols=3)
meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_items = [('📅 작성일', '2026년 3월'), ('📋 버전', 'v1.0'), ('🏭 대상', '제조 공장 에너지·ESG 담당자')]
for i, (label, val) in enumerate(meta_items):
    cell = meta_tbl.cell(0, i)
    set_cell_bg(cell, 'F0F9FF')
    set_cell_border(cell, 'BFDBFE')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_lbl = p.add_run(f'{label}\n')
    set_run_style(run_lbl, bold=True, size=9, color=(0x1e, 0x40, 0xaf))
    run_val = p.add_run(val)
    set_run_style(run_val, size=9, color=(0x37, 0x41, 0x5B))
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 목차
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('📋 목차', 1)
toc_items = [
    ('1', '이 대시보드는 무엇인가요?', '2'),
    ('2', '대시보드 목적', '2'),
    ('3', '사용 방법 (단계별)', '3'),
    ('4', '입력값 설명 — 확정값 vs 추측값', '4'),
    ('5', '결과 읽는 방법', '5'),
    ('6', '용어 설명 (초등학생도 이해할 수 있게!)', '6'),
    ('7', '자주 묻는 질문 (FAQ)', '7'),
]
toc_tbl = doc.add_table(rows=len(toc_items), cols=3)
toc_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (num, title, page) in enumerate(toc_items):
    cells = toc_tbl.rows[i].cells
    bg = 'EFF6FF' if i % 2 == 0 else 'F8FAFC'
    for cell in cells:
        set_cell_bg(cell, bg)
        set_cell_border(cell, 'DBEAFE')
    r0 = cells[0].paragraphs[0].add_run(f' {num}')
    set_run_style(r0, bold=True, size=10, color=(0x1e, 0x40, 0xaf))
    r1 = cells[1].paragraphs[0].add_run(title)
    set_run_style(r1, size=10)
    r2 = cells[2].paragraphs[0].add_run(f'{page}p')
    set_run_style(r2, size=10, color=(0x64, 0x74, 0x8B))
    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cell in cells:
        cell.paragraphs[0].paragraph_format.space_before = Pt(3)
        cell.paragraphs[0].paragraph_format.space_after = Pt(3)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 - 이 대시보드는 무엇인가요?
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('1️⃣  이 대시보드는 무엇인가요?', 1)

add_note_box(
    '💡 핵심 한 줄 요약\n'
    '"DR(수요반응) 발령이 오면 우리 공장이 얼마나 이익을 얻는지 미리 계산해주는 도구입니다."',
    bg='EFF6FF', border_color='2563EB'
)

add_heading('DR(수요반응)이란?', 2)
add_body('⚡ 여름철이나 겨울철에 전기가 부족해지면, 한국전력(한전)이 공장에 이렇게 말해요:', bullet=False)
add_note_box(
    '🏭 한전 → 공장  :  "지금 전기가 모자라요! 잠깐 전기를 좀 줄여주세요~"',
    bg='FEF9C3', border_color='F59E0B'
)
add_body('그러면 공장이 전기 사용량을 줄여주는 대신, 한전이 보상금(정산금)을 줍니다. 이 약속을 DR(수요반응, Demand Response)이라고 해요!', bullet=False)

add_heading('이 대시보드가 하는 일', 2)
items_s1 = [
    ('💰', 'DR 보상금(정산금) 미리 계산', '한전이 얼마나 줄지 예상해봅니다'),
    ('💡', '전기요금 절감액 확인', '전기를 줄이면 요금도 줄어들어요'),
    ('🌿', '탄소 감축량 계산', '환경에도 좋은 효과를 숫자로 확인해요'),
    ('📊', '한눈에 비교', '차트와 카드로 쉽게 결과를 볼 수 있어요'),
]
feat_tbl = doc.add_table(rows=2, cols=2)
feat_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for idx, (icon, title, desc) in enumerate(items_s1):
    row = idx // 2
    col = idx % 2
    cell = feat_tbl.cell(row, col)
    set_cell_bg(cell, 'F0F9FF')
    set_cell_border(cell, 'BFDBFE')
    p = cell.paragraphs[0]
    r1 = p.add_run(f'{icon} {title}\n')
    set_run_style(r1, bold=True, size=10.5, color=(0x1e, 0x40, 0xaf))
    r2 = p.add_run(desc)
    set_run_style(r2, size=9.5, color=(0x37, 0x41, 0x5B))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.1)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 - 대시보드 목적
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('2️⃣  대시보드 목적', 1)

add_heading('누가 사용하나요?', 2)
users_tbl = doc.add_table(rows=4, cols=2)
users_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
users_tbl.style = 'Table Grid'
headers = ['사용자', '활용 목적']
for i, h in enumerate(headers):
    cell = users_tbl.cell(0, i)
    set_cell_bg(cell, '1E40AF')
    set_cell_border(cell, '1E40AF')
    r = cell.paragraphs[0].add_run(h)
    set_run_style(r, bold=True, size=10, color=(255, 255, 255))
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

user_rows = [
    ('🏭 제조 공장 에너지 담당자', 'DR 발령 시 예상 이익 사전 계산, DR 참여 여부 판단'),
    ('🌿 ESG 담당자 (대기업)', '탄소 감축 효과를 숫자로 측정 → ESG 보고서 작성에 활용'),
    ('📊 경영진 보고용', '연간 DR 수익 전망치 확인 → 투자 의사결정 지원'),
]
for i, (user, purpose) in enumerate(user_rows):
    cells = users_tbl.rows[i+1].cells
    bg = 'EFF6FF' if i % 2 == 0 else 'FFFFFF'
    for cell in cells:
        set_cell_bg(cell, bg)
        set_cell_border(cell, 'DBEAFE')
    r0 = cells[0].paragraphs[0].add_run(user)
    set_run_style(r0, bold=True, size=10)
    r1 = cells[1].paragraphs[0].add_run(purpose)
    set_run_style(r1, size=10)
    for cell in cells:
        cell.paragraphs[0].paragraph_format.space_before = Pt(3)
        cell.paragraphs[0].paragraph_format.space_after = Pt(3)
        cell.paragraphs[0].paragraph_format.left_indent = Inches(0.05)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 - 사용 방법
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('3️⃣  사용 방법 (단계별)', 1)

steps = [
    (
        'Step 1', '데이터 입력 방식 선택',
        [
            ('✏️ 직접 입력', '숫자를 직접 타이핑합니다. 데이터 파일이 없어도 OK!\n평균 전력 사용량(kWh)만 알면 됩니다.'),
            ('📁 CSV 파일 업로드', '과거 전력 데이터 파일을 업로드합니다. 더 정확한 결과를 얻을 수 있어요.\n(파일 형식: 날짜, 시간, 평균(kWh) 컬럼 필요)'),
        ]
    ),
    (
        'Step 2', 'DR 발령 조건 입력',
        [
            ('📅 발령 월', 'DR이 발생하는 달을 선택합니다. (여름·겨울에 많이 발령됩니다)'),
            ('⏰ 감축 시작/종료 시간', '한전이 전기를 줄여달라고 요청하는 시간 범위입니다'),
            ('💰 DR 정산 단가', '전력 1kWh를 줄일 때 한전이 주는 보상금 (한전 계약서 확인)'),
            ('📉 감축 목표율', '평소보다 몇 % 줄일지 설정합니다'),
            ('🔢 월 DR 발령 횟수', '한 달에 몇 번 DR이 발령되는지'),
            ('📆 연간 DR 발령 월 수', '1년 중 DR이 발령되는 달 수'),
        ]
    ),
    (
        'Step 3', '공장 전력 정보 입력 (직접 입력 방식일 때)',
        [
            ('⚡ 시간당 평균 전력 사용량', 'DR 발령 시간대에 공장이 평균적으로 사용하는 전력량(kWh)을 입력합니다'),
        ]
    ),
    (
        'Step 4', '시뮬레이션 실행',
        [
            ('▶️ 버튼 클릭', '"🚀 DR 시뮬레이션 실행" 버튼을 클릭하면 바로 결과가 나타납니다!'),
        ]
    ),
    (
        'Step 5', '결과 확인',
        [
            ('📊 상단 4개 카드', '감축량 / DR 정산금 / 전기요금 절감 / 탄소 감축 수치'),
            ('📈 진행 바', 'CBL 대비 감축 목표를 시각적으로 표시'),
            ('🧮 정산 상세', '계산 과정 상세 내역 (어떻게 계산했는지 확인 가능)'),
            ('📉 차트', 'CBL vs 감축 목표 비교 그래프'),
            ('🌿 ESG 탄소', '연간 탄소 감축 효과 및 탄소 비용 절감액'),
        ]
    ),
]

for step_title, step_name, sub_items in steps:
    # Step 제목 배너
    step_tbl = doc.add_table(rows=1, cols=1)
    step_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    step_cell = step_tbl.cell(0, 0)
    set_cell_bg(step_cell, '1E3A8A')
    p = step_cell.paragraphs[0]
    r1 = p.add_run(f'  {step_title}  |  ')
    set_run_style(r1, bold=True, size=11, color=(147, 197, 253))
    r2 = p.add_run(step_name)
    set_run_style(r2, bold=True, size=11, color=(255, 255, 255))
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    # 서브 아이템
    for label, desc in sub_items:
        sub_tbl = doc.add_table(rows=1, cols=2)
        sub_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        c0 = sub_tbl.cell(0, 0)
        c1 = sub_tbl.cell(0, 1)
        set_cell_bg(c0, 'DBEAFE')
        set_cell_bg(c1, 'F8FAFC')
        set_cell_border(c0, 'BFDBFE')
        set_cell_border(c1, 'E2E8F0')
        # 첫 번째 열 너비 고정
        c0._tc.tcPr.append(OxmlElement('w:tcW'))
        tcW = c0._tc.tcPr.find(qn('w:tcW'))
        if tcW is None:
            tcW = OxmlElement('w:tcW')
            c0._tc.tcPr.append(tcW)
        tcW.set(qn('w:w'), '2000')
        tcW.set(qn('w:type'), 'dxa')

        r0 = c0.paragraphs[0].add_run(label)
        set_run_style(r0, bold=True, size=10, color=(0x1e, 0x40, 0xaf))
        r1 = c1.paragraphs[0].add_run(desc)
        set_run_style(r1, size=10)
        for c in [c0, c1]:
            c.paragraphs[0].paragraph_format.space_before = Pt(3)
            c.paragraphs[0].paragraph_format.space_after = Pt(3)
            c.paragraphs[0].paragraph_format.left_indent = Inches(0.05)

    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 - 확정값 vs 추측값
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('4️⃣  입력값 설명 — 확정값 vs 추측값', 1)

add_note_box(
    '⚠️  중요한 구분!\n'
    '• 확정값: 법이나 한전 공식 기준으로 정해진 값 → 바꾸지 않아도 됩니다\n'
    '• 추측값: 회사마다 다를 수 있는 값 → 실제 계약서나 담당자 확인이 필요합니다',
    bg='FFF7ED', border_color='F59E0B'
)

add_heading('✅  확정된 값 (공식 기준)', 2)

confirmed_data = [
    ('CBL 계산 방식', '직전 10 평일 중 상위 4일 평균', '한전 DR 공식 계산 기준'),
    ('전력량 요금 단가', '한전 고압A 선택I 요금표', '한전 고시 요금 (엑셀 파일 자동 로드)'),
    ('탄소 배출계수', '0.4781 kgCO₂/kWh', '환경부 고시 2022년 기준'),
    ('부가가치세', '10%', '세법 기준'),
    ('전력산업기반기금', '3.7%', '전기사업법 기준'),
]

conf_tbl = doc.add_table(rows=len(confirmed_data)+1, cols=3)
conf_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
conf_tbl.style = 'Table Grid'
conf_headers = ['항목', '값', '근거']
for i, h in enumerate(conf_headers):
    cell = conf_tbl.cell(0, i)
    set_cell_bg(cell, '166534')  # 초록 계열
    set_cell_border(cell, '166534')
    r = cell.paragraphs[0].add_run(h)
    set_run_style(r, bold=True, size=10, color=(255, 255, 255))
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, (item, val, reason) in enumerate(confirmed_data):
    cells = conf_tbl.rows[i+1].cells
    bg = 'F0FDF4' if i % 2 == 0 else 'DCFCE7'
    for cell in cells:
        set_cell_bg(cell, bg)
        set_cell_border(cell, 'BBF7D0')
    texts = [item, val, reason]
    for j, text in enumerate(texts):
        r = cells[j].paragraphs[0].add_run(text)
        set_run_style(r, bold=(j == 0), size=10, color=(0x14, 0x53, 0x2D) if j == 0 else None)
        cells[j].paragraphs[0].paragraph_format.space_before = Pt(3)
        cells[j].paragraphs[0].paragraph_format.space_after = Pt(3)
        cells[j].paragraphs[0].paragraph_format.left_indent = Inches(0.05)

doc.add_paragraph()

add_heading('⚠️  추측값 (사용자 확인 필요)', 2)

estimated_data = [
    ('DR 정산 단가', '300원/kWh (기본값)', '실제 한전 DR 계약서 확인 필요', '한전 DR 담당자 문의'),
    ('감축 목표율', '15% (기본값)', '공장마다 다름', '현장 에너지 담당자 협의'),
    ('월 DR 발령 횟수', '2회 (기본값)', '계약 조건에 따라 다름', '한전 DR 계약 조건 확인'),
    ('연간 DR 발령 월 수', '4개월 (기본값)', '연도·지역별 다름', '한전 DR 발령 이력 확인'),
    ('탄소 가격', '10,000원/톤', '시장 가격 참고값', '한국 ETS 시장 가격 참고'),
]

est_tbl = doc.add_table(rows=len(estimated_data)+1, cols=4)
est_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
est_tbl.style = 'Table Grid'
est_headers = ['항목', '기본값', '왜 추측인가요?', '확인 방법']
for i, h in enumerate(est_headers):
    cell = est_tbl.cell(0, i)
    set_cell_bg(cell, 'C2410C')  # 주황 계열
    set_cell_border(cell, 'C2410C')
    r = cell.paragraphs[0].add_run(h)
    set_run_style(r, bold=True, size=10, color=(255, 255, 255))
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, row_data in enumerate(estimated_data):
    cells = est_tbl.rows[i+1].cells
    bg = 'FFF7ED' if i % 2 == 0 else 'FFEDD5'
    for cell in cells:
        set_cell_bg(cell, bg)
        set_cell_border(cell, 'FED7AA')
    for j, text in enumerate(row_data):
        r = cells[j].paragraphs[0].add_run(text)
        set_run_style(r, bold=(j == 0), size=9.5, color=(0x7C, 0x2D, 0x12) if j == 0 else None)
        cells[j].paragraphs[0].paragraph_format.space_before = Pt(3)
        cells[j].paragraphs[0].paragraph_format.space_after = Pt(3)
        cells[j].paragraphs[0].paragraph_format.left_indent = Inches(0.05)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 - 결과 읽는 방법
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('5️⃣  결과 읽는 방법', 1)

add_heading('📊 상단 4개 결과 카드', 2)

result_cards = [
    ('⚡ 1회 총 감축량', 'kWh', '이번 DR 발령에서 줄이는 전력량 합계\n(= 시간당 감축량 × DR 발령 시간 수)'),
    ('💰 DR 정산금', '원', '한전이 공장에 지급하는 보상금\n(= 총 감축량 × DR 정산 단가)'),
    ('💡 전기요금 절감', '원', '전기를 덜 써서 아끼는 전기요금\n(= 시간당 감축량 × 한전 요금 단가)'),
    ('🌿 탄소 감축', 'kgCO₂', '줄어드는 이산화탄소 배출량\n(= 총 감축량 × 0.4781 kgCO₂/kWh)'),
]
rc_tbl = doc.add_table(rows=2, cols=2)
rc_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
card_colors = ['DBEAFE', 'DCFCE7', 'FEF9C3', 'D1FAE5']
border_colors = ['3B82F6', '22C55E', 'F59E0B', '10B981']
for idx, (title, unit, desc) in enumerate(result_cards):
    row = idx // 2
    col = idx % 2
    cell = rc_tbl.cell(row, col)
    set_cell_bg(cell, card_colors[idx])
    set_cell_border(cell, border_colors[idx])
    p = cell.paragraphs[0]
    r1 = p.add_run(f'{title}\n')
    set_run_style(r1, bold=True, size=11, color=(0x1e, 0x40, 0xaf))
    r2 = p.add_run(f'단위: {unit}\n')
    set_run_style(r2, size=9, color=(0x64, 0x74, 0x8B))
    r3 = p.add_run(desc)
    set_run_style(r3, size=9.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.1)

doc.add_paragraph()

add_heading('🧮 총 이익 계산 공식', 2)
add_note_box(
    '총 예상 이익 (1회) = DR 정산금 + 전기요금 절감액\n\n'
    '연간 DR 총 이익 = 총 예상 이익 (1회) × 월 DR 발령 횟수 × 연간 DR 발령 월 수',
    bg='F0F9FF', border_color='0EA5E9'
)

add_heading('📈 차트 읽는 법', 2)
chart_items = [
    ('파란 막대 (CBL)', '평소 전력 사용량 기준선 — 이것을 "기준값"으로 사용해요'),
    ('초록 막대 (감축 목표)', 'DR 발령 시 줄인 후의 목표 전력량'),
    ('두 막대의 차이', '이 부분이 실제로 "감축된 전력"이에요!'),
]
for icon_label, desc in chart_items:
    add_body(f'• {icon_label}: {desc}')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6 - 용어 설명
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('6️⃣  용어 설명 (초등학생도 이해할 수 있게!)', 1)

terms = [
    ('DR (수요반응)', 'Demand Response', '한전이 "전기 좀 줄여주세요~" 하고 공장에 요청하는 것이에요. 공장이 도와주면 한전이 보상금을 줘요!'),
    ('발령', '-', '"한전이 공식적으로 요청을 보내는 것"입니다. 마치 선생님이 "지금 조용히 하세요!"라고 말하는 것처럼요.'),
    ('CBL (기준 부하)', 'Customer Baseline Load', '평소에 공장이 사용하는 전력의 기준 값입니다. 직전 10 평일 중 전력 사용량이 많았던 상위 4일의 평균이에요.'),
    ('정산금', '-', '전기를 줄여준 것에 대해 한전이 지급하는 보상금입니다. 전력 감축량 × 정산 단가로 계산해요.'),
    ('배출계수', '0.4781 kgCO₂/kWh', '전기를 1kWh 쓸 때 나오는 이산화탄소(CO₂)의 양이에요. 환경부가 정해준 공식 수치입니다.'),
    ('ESG', 'E(환경) S(사회) G(지배구조)', '기업이 환경·사회·윤리적으로 잘 운영되고 있는지 보여주는 지표예요. 탄소 감축이 E(환경) 부분에 해당합니다.'),
    ('kWh', 'kilowatt-hour', '전력 사용량의 단위예요. 1kWh는 1,000W짜리 전자기기를 1시간 쓴 것과 같아요.'),
    ('피크', 'Peak', '전력 사용이 가장 많은 시간대를 말해요. 주로 여름 낮 시간(14~17시), 겨울 저녁(18~21시)이에요.'),
    ('경부하 / 중간부하 / 최대부하', '-', '전력 사용량에 따라 시간대를 나눈 것이에요.\n• 경부하: 전기를 적게 쓰는 시간 (새벽 등)\n• 중간부하: 보통 쓰는 시간\n• 최대부하: 가장 많이 쓰는 시간 (요금이 가장 비쌈!)'),
    ('한전 (KEPCO)', '한국전력공사', '우리나라 전기를 관리하는 공공기관이에요. DR 계약을 맺고 정산금을 지급합니다.'),
    ('ETS', '배출권거래제 (Emissions Trading Scheme)', '기업이 온실가스를 얼마나 배출할 수 있는지 권리를 사고파는 시장이에요. 탄소 가격(10,000원/톤)이 여기서 결정돼요.'),
]

term_tbl = doc.add_table(rows=len(terms)+1, cols=3)
term_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
term_tbl.style = 'Table Grid'
term_headers = ['용어', '영어/약어', '쉬운 설명']
for i, h in enumerate(term_headers):
    cell = term_tbl.cell(0, i)
    set_cell_bg(cell, '4338CA')
    set_cell_border(cell, '4338CA')
    r = cell.paragraphs[0].add_run(h)
    set_run_style(r, bold=True, size=10, color=(255, 255, 255))
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, (term, eng, explanation) in enumerate(terms):
    cells = term_tbl.rows[i+1].cells
    bg = 'EEF2FF' if i % 2 == 0 else 'F5F3FF'
    for cell in cells:
        set_cell_bg(cell, bg)
        set_cell_border(cell, 'C7D2FE')
    r0 = cells[0].paragraphs[0].add_run(term)
    set_run_style(r0, bold=True, size=10, color=(0x31, 0x27, 0x9A))
    r1 = cells[1].paragraphs[0].add_run(eng)
    set_run_style(r1, size=9, color=(0x64, 0x74, 0x8B))
    r2 = cells[2].paragraphs[0].add_run(explanation)
    set_run_style(r2, size=9.5)
    for cell in cells:
        cell.paragraphs[0].paragraph_format.space_before = Pt(3)
        cell.paragraphs[0].paragraph_format.space_after = Pt(3)
        cell.paragraphs[0].paragraph_format.left_indent = Inches(0.05)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7 - FAQ
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('7️⃣  자주 묻는 질문 (FAQ)', 1)

faqs = [
    (
        'Q1. 데이터가 없어도 사용할 수 있나요?',
        '네! "✏️ 직접 입력" 방식을 선택하면 평균 전력 사용량(kWh)만 알면 됩니다. 과거 데이터 파일이 없어도 충분히 시뮬레이션할 수 있어요.'
    ),
    (
        'Q2. CSV 파일 형식이 어떻게 되나요?',
        '아래 컬럼이 필요합니다:\n• 날짜: YYYYMMDD 형식 (예: 20240715)\n• 시간: 0~23 숫자\n• 평균: kWh 단위의 시간당 평균 전력 사용량\nUTF-8 또는 EUC-KR 인코딩을 모두 지원합니다.'
    ),
    (
        'Q3. DR 정산 단가를 모르면 어떻게 하나요?',
        '한전 DR 담당자에게 문의하거나 DR 계약서를 확인하세요. 계약서가 없을 경우 임시로 기본값 300원/kWh를 사용해도 됩니다. (단, 실제 값과 다를 수 있습니다)'
    ),
    (
        'Q4. 결과가 실제와 다를 수 있나요?',
        '네. DR 정산 단가, 감축 목표율, 월 DR 발령 횟수 등 추측값에 따라 결과가 달라집니다. 실제 한전 DR 계약 조건을 입력할수록 결과가 정확해집니다.'
    ),
    (
        'Q5. CBL이 너무 낮게 나와요. 왜 그런가요?',
        'CBL은 직전 10 평일의 상위 4일 평균으로 계산됩니다. CSV 파일의 해당 월 데이터가 부족하면 CBL 계산이 정확하지 않을 수 있어요. 최소 10개 평일 데이터가 있는 달을 선택하세요.'
    ),
    (
        'Q6. 탄소 배출계수(0.4781)는 언제 업데이트되나요?',
        '환경부가 매년 고시합니다. 현재 대시보드는 2022년 기준값을 사용합니다. 최신 값은 환경부 홈페이지에서 확인할 수 있습니다.'
    ),
]

for q, a in faqs:
    # Q 박스
    q_tbl = doc.add_table(rows=1, cols=1)
    q_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    q_cell = q_tbl.cell(0, 0)
    set_cell_bg(q_cell, '1E40AF')
    p = q_cell.paragraphs[0]
    r = p.add_run(f'❓ {q}')
    set_run_style(r, bold=True, size=10.5, color=(255, 255, 255))
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)

    # A 박스
    a_tbl = doc.add_table(rows=1, cols=1)
    a_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    a_cell = a_tbl.cell(0, 0)
    set_cell_bg(a_cell, 'EFF6FF')
    p = a_cell.paragraphs[0]
    r = p.add_run(f'✅ {a}')
    set_run_style(r, size=10)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)

    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 대시보드 구조도
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading('📐 대시보드 구조도', 1)

add_note_box(
    '아래 구조도는 DR 시뮬레이션 대시보드의 전체 흐름을 보여줍니다.\n'
    '입력값이 어떻게 처리되어 결과로 나오는지 한눈에 볼 수 있어요.',
    bg='EFF6FF', border_color='2563EB'
)

add_heading('전체 흐름도', 2)

flow_data = [
    ('1️⃣ 입력', '데이터 입력 방식 선택', ['✏️ 직접 입력 (평균 kWh)', '📁 CSV 파일 업로드 (과거 데이터)']),
    ('2️⃣ 조건', 'DR 발령 조건 설정', ['📅 발령 월 선택', '⏰ 감축 시간 범위', '💰 DR 정산 단가', '📉 감축 목표율 (%)', '🔢 월 발령 횟수 / 연간 발령 월 수']),
    ('3️⃣ 계산', '시뮬레이션 엔진', ['CBL 산출 (10평일 상위 4일 평균)', '감축량 계산 (CBL × 감축 목표율)', 'DR 정산금 계산', '전기요금 절감 계산', '탄소 감축 계산']),
    ('4️⃣ 결과', '결과 화면 출력', ['📊 4개 결과 카드 (감축량/정산금/절감/탄소)', '📈 CBL vs 감축 목표 차트', '🧮 정산 상세 내역', '🌿 ESG 연간 탄소 감축 효과']),
]

for stage, title, items in flow_data:
    flow_tbl = doc.add_table(rows=1, cols=2)
    flow_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    c0 = flow_tbl.cell(0, 0)
    c1 = flow_tbl.cell(0, 1)
    set_cell_bg(c0, '1E3A8A')
    set_cell_bg(c1, 'F0F9FF')
    set_cell_border(c0, '1E40AF')
    set_cell_border(c1, 'BFDBFE')

    p0 = c0.paragraphs[0]
    r0a = p0.add_run(f'{stage}\n')
    set_run_style(r0a, bold=True, size=13, color=(147, 197, 253))
    r0b = p0.add_run(title)
    set_run_style(r0b, bold=True, size=10, color=(255, 255, 255))
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(8)
    p0.paragraph_format.space_after = Pt(8)

    for i, item in enumerate(items):
        if i == 0:
            p1 = c1.paragraphs[0]
        else:
            p1 = c1.add_paragraph()
        r1 = p1.add_run(f'  • {item}')
        set_run_style(r1, size=10)
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)

    # 화살표 (다음 단계가 있을 경우)
    if stage != '4️⃣ 결과':
        arrow_p = doc.add_paragraph()
        arrow_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = arrow_p.add_run('▼')
        set_run_style(r, bold=True, size=14, color=(0x1e, 0x40, 0xaf))
        arrow_p.paragraph_format.space_before = Pt(0)
        arrow_p.paragraph_format.space_after = Pt(0)
    else:
        doc.add_paragraph()

add_heading('데이터 흐름 상세', 2)

flow_detail = [
    ('입력값', '→', 'CBL 계산', '직전 10 평일 데이터 또는 직접 입력한 평균 kWh 사용'),
    ('CBL', '→', '감축량 산출', 'CBL × (1 - 감축 목표율 %) = 감축 후 목표 전력\n감축량 = CBL - 목표 전력'),
    ('감축량', '→', 'DR 정산금', '총 감축량(kWh) × DR 정산 단가(원/kWh)'),
    ('감축량', '→', '전기요금 절감', '시간당 감축량 × 한전 요금 단가 (시간대별 다름)'),
    ('감축량', '→', '탄소 감축', '총 감축량 × 0.4781 kgCO₂/kWh (환경부 고시)'),
    ('모든 결과', '→', '연간 환산', '× 월 DR 발령 횟수 × 연간 DR 발령 월 수'),
]

det_tbl = doc.add_table(rows=len(flow_detail)+1, cols=4)
det_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
det_tbl.style = 'Table Grid'
det_headers = ['시작', '', '결과', '계산 방법']
for i, h in enumerate(det_headers):
    cell = det_tbl.cell(0, i)
    set_cell_bg(cell, '374151')
    r = cell.paragraphs[0].add_run(h)
    set_run_style(r, bold=True, size=10, color=(255, 255, 255))
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, (start, arrow, result, method) in enumerate(flow_detail):
    cells = det_tbl.rows[i+1].cells
    bg = 'F9FAFB' if i % 2 == 0 else 'FFFFFF'
    for cell in cells:
        set_cell_bg(cell, bg)
        set_cell_border(cell, 'E5E7EB')
    texts_styles = [
        (start, (0x1e, 0x40, 0xaf), True),
        (arrow, (0x6B, 0x72, 0x80), False),
        (result, (0x05, 0x60, 0x32), True),
        (method, None, False),
    ]
    for j, (text, color, bold) in enumerate(texts_styles):
        r = cells[j].paragraphs[0].add_run(text)
        set_run_style(r, bold=bold, size=9.5, color=color)
        cells[j].paragraphs[0].paragraph_format.space_before = Pt(3)
        cells[j].paragraphs[0].paragraph_format.space_after = Pt(3)
        cells[j].paragraphs[0].paragraph_format.left_indent = Inches(0.05)

doc.add_paragraph()

# ─── 푸터 ──────────────────────────────────────────────────────────────────
add_note_box(
    '📌  이 문서는 DR 시뮬레이션 대시보드(DR_app.py) v1.0 기준으로 작성되었습니다.\n'
    '문의사항이 있으면 에너지 담당자 또는 IT 담당자에게 연락하세요.',
    bg='F8FAFC', border_color='94A3B8'
)

# ─── 저장 ──────────────────────────────────────────────────────────────────
output_path = r'C:\Users\Admin\hipython\ml\2차 프로젝트\DR대시보드_설명서.docx'
doc.save(output_path)
print(f'✅ 문서 저장 완료: {output_path}')
